import argparse
import re
import sys
from datetime import date, datetime
from pathlib import Path

from docx import Document
from openpyxl import Workbook, load_workbook


COLUMN_ALIASES = {
    "party_name": [
        "party_name",
        "party",
        "counterparty",
        "往来单位",
        "单位名称",
        "客户名称",
        "供应商名称",
        "对方单位",
        "单位",
        "客户",
        "供应商",
        "客商名称",
    ],
    "amount": [
        "amount",
        "balance",
        "余额",
        "期末余额",
        "本币余额",
        "金额",
        "对账金额",
        "确认金额",
    ],
    "balance_date": [
        "balance_date",
        "date",
        "截止日期",
        "结算日期",
        "余额日期",
        "对账日期",
        "日期",
    ],
    "address": ["address", "地址", "联系地址"],
    "contact": ["contact", "联系人", "联系方式"],
    "currency": ["currency", "币种", "币别"],
    "remarks": ["remarks", "备注", "说明"],
}


def normalize_header(value):
    if value is None:
        return ""
    text = str(value).strip().lower()
    text = text.replace("（", "(").replace("）", ")")
    text = re.sub(r"\s+", "", text)
    text = re.sub(r"[()（）\[\]【】:%/\\-]", "", text)
    return text


def find_header_index(headers, aliases):
    alias_norms = [normalize_header(a) for a in aliases]
    for alias in alias_norms:
        if alias in headers:
            return headers.index(alias)
    for idx, header in enumerate(headers):
        for alias in alias_norms:
            if alias and alias in header:
                return idx
    return None


def format_amount(value):
    if value is None:
        return ""
    try:
        text = str(value).replace(",", "").strip()
        num = float(text)
        return f"{num:,.2f}"
    except (TypeError, ValueError):
        return str(value).strip()


def format_date(value):
    if value is None:
        return ""
    if isinstance(value, (datetime, date)):
        return value.strftime("%Y-%m-%d")
    return str(value).strip()


def sanitize_filename(name):
    name = name.strip()
    name = re.sub(r'[<>:"/\\|?*]', "_", name)
    name = re.sub(r"\s+", " ", name)
    return name[:120] if name else ""


def build_default_template():
    doc = Document()
    doc.add_heading("Confirmation Letter", level=1)
    doc.add_paragraph("To: {{party_name}}")
    doc.add_paragraph("Address: {{address}}")
    doc.add_paragraph("")
    doc.add_paragraph(
        "As of {{balance_date}}, our records show a balance of {{currency}} {{amount_formatted}}."
    )
    doc.add_paragraph("Please confirm the balance and reply with your confirmation.")
    doc.add_paragraph("")
    doc.add_paragraph("Contact: {{contact}}")
    doc.add_paragraph("Remarks: {{remarks}}")
    return doc


def make_template(path):
    doc = build_default_template()
    doc.save(path)


def replace_placeholders_in_paragraph(paragraph, data):
    if "{{" not in paragraph.text:
        return
    replaced = False
    for run in paragraph.runs:
        text = run.text
        if "{{" not in text:
            continue
        for key, val in data.items():
            text = text.replace(f"{{{{{key}}}}}", str(val))
        if text != run.text:
            run.text = text
            replaced = True
    if "{{" in paragraph.text and not replaced:
        text = paragraph.text
        for key, val in data.items():
            text = text.replace(f"{{{{{key}}}}}", str(val))
        paragraph.text = text


def fill_placeholders(doc, data):
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, data)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_placeholders_in_paragraph(paragraph, data)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholders_in_paragraph(paragraph, data)
        for paragraph in section.footer.paragraphs:
            replace_placeholders_in_paragraph(paragraph, data)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholders_in_paragraph(paragraph, data)


def read_rows(ws):
    required = ["party_name", "amount", "balance_date"]
    optional = ["address", "contact", "currency", "remarks"]
    header_row = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))
    headers = [normalize_header(value) for value in header_row]

    indices = {}
    for key in required + optional:
        idx = find_header_index(headers, COLUMN_ALIASES.get(key, [key]))
        if key in required and idx is None:
            raise ValueError(f"Missing required column: {key}")
        indices[key] = idx

    rows = []
    for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        if all(v is None for v in row):
            continue
        data = {}
        for key in required + optional:
            idx = indices.get(key)
            data[key] = row[idx] if idx is not None else None
        if not data["party_name"]:
            raise ValueError(f"Row {row_idx} missing party_name.")
        rows.append(data)
    return rows


def write_index(output_dir, rows):
    wb = Workbook()
    ws = wb.active
    ws.title = "Index"
    ws.append(["party_name", "amount", "balance_date", "file"])
    for row in rows:
        ws.append(row)
    wb.save(output_dir / "index.xlsx")


def main():
    parser = argparse.ArgumentParser(description="Generate confirmation letters from Excel.")
    parser.add_argument("--input", default="input.xlsx", help="Input .xlsx (default: input.xlsx).")
    parser.add_argument("--output", default="output", help="Output folder (default: output).")
    parser.add_argument("--sheet", default="", help="Sheet name (default: first sheet).")
    parser.add_argument("--template", default="", help="Template .docx with {{placeholders}}.")
    parser.add_argument("--make-template", default="", help="Create a sample template and exit.")
    args = parser.parse_args()

    if args.make_template:
        make_template(args.make_template)
        print(f"Template saved: {args.make_template}")
        return 0

    try:
        wb = load_workbook(args.input, data_only=True)
        ws = wb[args.sheet] if args.sheet else wb.active
        rows = read_rows(ws)
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1

    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)

    index_rows = []
    for idx, data in enumerate(rows, start=1):
        party_name = str(data["party_name"]).strip()
        amount = data["amount"]
        balance_date = data["balance_date"]

        data = {
            "party_name": party_name,
            "address": data.get("address") or "",
            "contact": data.get("contact") or "",
            "currency": data.get("currency") or "",
            "remarks": data.get("remarks") or "",
            "amount": amount if amount is not None else "",
            "amount_formatted": format_amount(amount),
            "balance_date": format_date(balance_date),
        }

        if args.template:
            doc = Document(args.template)
        else:
            doc = build_default_template()
        fill_placeholders(doc, data)

        safe_name = sanitize_filename(party_name) or f"party_{idx}"
        filename = f"{safe_name}.docx"
        doc.save(output_dir / filename)
        index_rows.append([party_name, amount, format_date(balance_date), filename])

    write_index(output_dir, index_rows)
    print(f"Generated {len(index_rows)} letter(s) in {output_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
